"""
export_docx.py — Faithful Word DOCX conversion of the PDF carnet.

Uses export_content.py to extract maps, encoded text, annexes, and fonts.
"""
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.pdf_helpers import get_theme_label


def _accent_rgb(is_sol):
    return RGBColor(0xb4, 0x4a, 0x2b) if is_sol else RGBColor(0x2d, 0x5a, 0x8e)


def _add_heading(doc, text, level, is_sol):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _accent_rgb(is_sol)
    return h


def _add_map_image(doc, map_image, caption='Carte IGN'):
    if map_image is None:
        return
    buf = BytesIO()
    map_image.save(buf, format='PNG')
    buf.seek(0)
    try:
        p = doc.add_picture(buf, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Caption
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    except Exception as e:
        doc.add_paragraph(f"[Carte non disponible : {e}]")


def _add_gilwell_image(doc, gilwell_svg):
    """Render Gilwell SVG to PNG via reportlab/cairosvg or fallback to text."""
    if not gilwell_svg:
        return
    try:
        import cairosvg
        png_data = cairosvg.svg2png(bytestring=gilwell_svg.encode('utf-8'))
        doc.add_picture(BytesIO(png_data), width=Inches(3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except ImportError:
        # Fallback: extract azimut text from SVG (basic)
        import re
        texts = re.findall(r'>(\d+°[^<]+)<', gilwell_svg)
        p = doc.add_paragraph()
        p.add_run("Relevé Gilwell :\n").bold = True
        for t in texts:
            p.add_run(f"  • {t}\n")


def _add_polybe_table(doc, grid_rows, accent_rgb):
    doc.add_paragraph("CARRÉ DE POLYBE").runs[0].bold = True
    table = doc.add_table(rows=7, cols=7)
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = ''
    for col in range(1, 7):
        hdr.cells[col].text = str(col)
        run = hdr.cells[col].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = accent_rgb
    # Data rows
    for row_i, cells in enumerate(grid_rows):
        row = table.rows[row_i + 1]
        row.cells[0].text = str(row_i + 1)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for col_i, (char, _coord) in enumerate(cells):
            row.cells[col_i + 1].text = char
    doc.add_paragraph("Lecture : Ligne (1er chiffre) puis Colonne (2ème chiffre). Ex: 21=G").style = 'Caption'


def _add_vigenere_note(doc, key, accent_rgb):
    p = doc.add_paragraph()
    p.add_run("CLÉ VIGENÈRE : ").bold = True
    key_run = p.add_run(key)
    key_run.bold = True
    key_run.font.color.rgb = accent_rgb
    key_run.font.size = Pt(14)
    doc.add_paragraph("(Voir table de Vigenère en annexe — Ligne = lettre de la clé, Colonne = lettre du message clair)")


def _add_morse_table(doc, morse_data):
    p = doc.add_paragraph("TABLE MORSE")
    p.runs[0].bold = True
    table = doc.add_table(rows=1 + len(morse_data) // 4 + 1, cols=8)
    table.style = 'Table Grid'
    items = list(morse_data.items())
    for i, (char, code) in enumerate(items):
        row = i // 4
        col = (i % 4) * 2
        try:
            table.rows[row + 1].cells[col].text = char
            table.rows[row + 1].cells[col + 1].text = code
        except Exception:
            pass


def _add_maritime_table(doc, nato):
    p = doc.add_paragraph("CODE NATO MARITIME")
    p.runs[0].bold = True
    cols_n = 4
    items = list(nato.items())
    table = doc.add_table(rows=len(items) // cols_n + 1, cols=cols_n * 2)
    table.style = 'Table Grid'
    for i, (char, word) in enumerate(items):
        row = i // cols_n
        col = (i % cols_n) * 2
        try:
            table.rows[row].cells[col].text = char
            table.rows[row].cells[col + 1].text = word
        except Exception:
            pass


def _render_docx(all_steps, is_sol):
    doc = Document()
    accent = _accent_rgb(is_sol)

    # Title
    title_str = get_theme_label('soluce_title', 'SOLUCE — CHEFS') if is_sol else get_theme_label('main_title', 'CARNET DE ROUTE')
    t = doc.add_heading(title_str, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = accent

    sub = doc.add_paragraph(f"{len(all_steps)} étapes · {'SOLUTION' if is_sol else 'PARTICIPANT'} · ScoutRaider Suite")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()

    for step in all_steps:
        # Step heading
        h = doc.add_heading(f"Étape {step['step_num']} — {step['label']}", level=1)
        for run in h.runs:
            run.font.color.rgb = accent
            run.font.size = Pt(14)

        # Map or Gilwell
        if step.get('map_image') is not None:
            _add_map_image(doc, step['map_image'], f"Carte – Étape {step['step_num']}")
        if step.get('gilwell_svg'):
            _add_gilwell_image(doc, step['gilwell_svg'])

        # Font install note
        font = step.get('font')
        if font and not is_sol and not os.path.isfile(font.get('file', '')):
            note = doc.add_paragraph()
            run = note.add_run(f"⚠ Police '{font['name']}' non installée. Installez le fichier depuis : modules/{step['module']}/assets/")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x92, 0x40, 0x0e)

        # Messages
        for idx, msg in enumerate(step.get('messages', []), 1):
            is_encoded = (not is_sol
                          and msg['type'] not in ('clair', 'texte_clair', 'drapeaux', 'visual')
                          and msg['encoded'] != msg['clair'])

            p = doc.add_paragraph()
            num_run = p.add_run(f"{idx}. ")
            num_run.bold = True
            num_run.font.color.rgb = accent

            if is_encoded:
                enc_run = p.add_run(msg['encoded'] + "\n")
                enc_run.font.name = 'Courier New'
                enc_run.font.size = Pt(11)
                hint_run = p.add_run(f"   → {msg['clair']}")
                hint_run.font.size = Pt(9)
                hint_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
            else:
                p.add_run(msg['clair'])

        # Annexes
        for annexe in step.get('annexes', []):
            atype = annexe['type']
            adata = annexe['data']
            doc.add_paragraph()
            if atype == 'polybe':
                _add_polybe_table(doc, adata, accent)
            elif atype == 'vigenere':
                _add_vigenere_note(doc, adata[0], accent)
            elif atype == 'morse':
                _add_morse_table(doc, adata)
            elif atype == 'maritime':
                _add_maritime_table(doc, adata)

        doc.add_paragraph()

    # Footer paragraph at end
    footer_p = doc.add_paragraph(f"ScoutRaider Suite · {get_theme_label('filename', 'Carnet')} · {'SOLUTION' if is_sol else 'Participant'}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_p.runs:
        footer_p.runs[0].font.size = Pt(9)
        footer_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    return doc


def _build_textual_carnet(orchestrator, path_plan, is_sol=False):
    """Legacy compat — used by export_odt.py."""
    from utils.export_content import extract_step_content
    steps = extract_step_content(orchestrator, path_plan, is_sol)
    blocks = []
    for step in steps:
        instructions = []
        for m in step['messages']:
            if is_sol or m['type'] in ('clair', 'texte_clair', 'drapeaux', 'visual'):
                instructions.append(m['clair'])
            else:
                instructions.append(m['encoded'])
        blocks.append({'step_num': step['step_num'], 'module': step['module'], 'instructions': instructions})
    return blocks


def export_docx(orchestrator, path_plan, output_dir=None, progress_callback=None, opts=None):
    if progress_callback: progress_callback("Récupération des POIs...", 20)

    import utils.pdf_helpers as ph
    ph.set_global_pois(ph.fetch_all_pois(orchestrator.segments))

    from utils.export_content import extract_step_content

    if progress_callback: progress_callback("Construction du contenu DOCX...", 35)
    steps     = extract_step_content(orchestrator, path_plan, is_sol=False)
    sol_steps = extract_step_content(orchestrator, path_plan, is_sol=True)

    base_name = get_theme_label('filename', 'Carnet')

    if progress_callback: progress_callback("Rendu DOCX Participant...", 55)
    out_path = os.path.join(output_dir, f"{base_name}.docx")
    _render_docx(steps, False).save(out_path)

    if progress_callback: progress_callback("Rendu DOCX Solution...", 80)
    sol_path = os.path.join(output_dir, f"{base_name}_SOLUCE.docx")
    _render_docx(sol_steps, True).save(sol_path)

    if progress_callback: progress_callback("Export DOCX terminé.", 100)
    return out_path, sol_path
